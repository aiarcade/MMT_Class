#!/usr/bin/env python3
"""
Script to create a properly formatted syllabus DOCX for Multimedia Technologies
Following the exact format from FUNDAMENTALS OF CYBER SECURITY.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def set_cell_borders(cell):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_bold_para(doc, text, size=12, align='left'):
    """Add a bold paragraph"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def create_table_with_borders(doc, rows, cols):
    """Create a table with borders"""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    return table

def main():
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # Header information
    p = doc.add_paragraph()
    run = p.add_run("Stream: Computer Science and Engineering")
    run.bold = True
    run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    run = p.add_run("Branch : CSD")
    run.bold = True
    run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    run = p.add_run("Semester : S6")
    run.bold = True
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    # Course Title
    p = doc.add_paragraph()
    run = p.add_run("MULTIMEDIA TECHNOLOGIES")
    run.bold = True
    run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Course Info Table
    table = create_table_with_borders(doc, 2, 9)
    
    # Header row
    headers = ['Course Code', 'L', 'T', 'P', 'R', 'Credits', 'Exam hours', 'Course type', 'CIE Marks', 'ESE Marks']
    header_cells = table.rows[0].cells
    header_cells[0].text = "Course Code"
    header_cells[1].text = "L"
    header_cells[2].text = "T"
    header_cells[3].text = "P"
    header_cells[4].text = "Credits"
    header_cells[5].text = "Exam hours"
    header_cells[6].text = "Course type"
    header_cells[7].text = "CIE Marks"
    header_cells[8].text = "ESE Marks"
    
    for cell in header_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data row
    data_cells = table.rows[1].cells
    data_cells[0].text = "CXT332"
    data_cells[1].text = "2"
    data_cells[2].text = "1"
    data_cells[3].text = "0"
    data_cells[4].text = "3"
    data_cells[5].text = "3 hrs"
    data_cells[6].text = "Theory"
    data_cells[7].text = "50"
    data_cells[8].text = "100"
    
    for cell in data_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 1. Preamble
    p = doc.add_paragraph()
    run = p.add_run("1. Preamble:")
    run.bold = True
    run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    p.add_run("This course provides a comprehensive understanding of multimedia systems, covering the fundamental concepts of digital media representation, compression techniques, and modern applications. The course enables students to apply contemporary theories of multimedia learning to understand data compression algorithms, industry standards, and cloud-based multimedia services. Through structured modules, students will gain practical knowledge in multimedia data handling, compression standards (JPEG, MPEG), content-based retrieval, and cloud computing for multimedia applications.")
    
    doc.add_paragraph()
    
    # 2. Prerequisites
    p = doc.add_paragraph()
    run = p.add_run("2. Prerequisites: ")
    run.bold = True
    run.font.size = Pt(11)
    p.add_run("Basic understanding of digital systems, data structures, and fundamentals of signal processing.")
    
    doc.add_paragraph()
    
    # 3. Course Objectives
    p = doc.add_paragraph()
    run = p.add_run("3. Course Objectives:")
    run.bold = True
    run.font.size = Pt(11)
    
    objectives = [
        "To understand multimedia data representations including graphics, audio, and video",
        "To master both lossless and lossy compression algorithms and their applications",
        "To apply knowledge of industry standards (JPEG, MPEG family) for multimedia compression",
        "To explore modern applications including content-based retrieval and cloud computing for multimedia"
    ]
    
    for i, obj in enumerate(objectives, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {obj}")
        p.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_paragraph()
    
    # Syllabus Table
    table = create_table_with_borders(doc, 6, 3)
    
    # Header
    header = table.rows[0].cells
    header[0].text = "Module Number"
    header[1].text = "Syllabus Description"
    header[2].text = "Contact Hours"
    for cell in header:
        set_cell_shading(cell, 'D9E2F3')
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Module 1
    row = table.rows[1].cells
    row[0].text = "1"
    row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row[1].text = """Multimedia Fundamentals & Data Representation:

Multimedia Basics- Multimedia, Hypermedia, WWW, Internet. Multimedia Software- Editing and Authoring Tools. Graphics/Image Data Types- Raster vs Vector, Bit depth, Resolution, Popular File Formats (BMP, GIF, PNG, TIFF, JPEG). Color Science- Human visual system, Color perception. Color Models- RGB, CMY/CMYK, HSV/HSL, YUV/YCbCr. Digital Audio- Digitization of sound, Sampling, Quantization, Nyquist theorem. MIDI- Musical Instrument Digital Interface. Digital Video- Video signals, Frame rates, Interlacing, Video formats."""
    row[2].text = "9 hours"
    row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Module 2
    row = table.rows[2].cells
    row[0].text = "2"
    row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row[1].text = """Compression Algorithms & Techniques:

Information Theory Basics- Entropy, Redundancy, Compression fundamentals. Run-Length Encoding (RLE)- Algorithm, Applications, Variants. Variable-Length Coding- Huffman coding, Shannon-Fano coding, Optimal prefix codes. Dictionary-Based Coding- LZ77, LZ78, LZW algorithms, Sliding window techniques. Arithmetic Coding- Principles. Lossy Compression Fundamentals- Distortion measures (MSE, PSNR), Rate-Distortion theory. Quantization Techniques- Scalar quantization, Vector quantization. Transform Coding."""
    row[2].text = "9 hours"
    row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Module 3
    row = table.rows[3].cells
    row[0].text = "3"
    row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row[1].text = """Multimedia Compression Standards:

Image Compression Standards- JPEG (Baseline, DCT-based compression, Quality settings). Video Compression- Motion estimation, Motion compensation, Block matching, MPEG-1 (Video bitstream structure, I/P/B frames), MPEG-2 (Interlaced video support), MPEG-4 (Object-based coding)."""
    row[2].text = "9 hours"
    row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Module 4
    row = table.rows[4].cells
    row[0].text = "4"
    row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row[1].text = """Advanced Applications & Cloud Multimedia:

Content-Based Image Retrieval (CBIR)- Feature extraction, Color histograms, Texture features. Similarity Measures & Indexing- Distance metrics, Indexing structures, Query processing. CBIR Case Study- CBIRD system, Practical implementation. Video Retrieval & Search- Video segmentation, Keyframe extraction, Video querying, Quantifying search results. Cloud Computing Overview- Cloud service models (IaaS, PaaS, SaaS), Deployment models (Public, Private, Hybrid). Multimedia Cloud Computing- Architecture, Components, Resource management. Cloud Media Services- Media sharing, Streaming services, CDN integration. Computation Offloading."""
    row[2].text = "9 hours"
    row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Total row
    row = table.rows[5].cells
    row[0].merge(row[1])
    row[0].text = "Total Hours"
    for para in row[0].paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in para.runs:
            run.bold = True
    row[2].text = "36 hours"
    for para in row[2].paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.bold = True
    
    doc.add_paragraph()
    
    # 4. Course Outcomes
    p = doc.add_paragraph()
    run = p.add_run("4. Course Outcomes")
    run.bold = True
    run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    p.add_run("At the end of the course students should be able to:")
    
    # CO Table
    table = create_table_with_borders(doc, 6, 3)
    
    header = table.rows[0].cells
    header[0].text = "CO"
    header[1].text = "Course Outcomes"
    header[2].text = "Blooms Taxonomy Level"
    for cell in header:
        set_cell_shading(cell, 'D9E2F3')
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    cos = [
        ("CO1", "Describe the basic concepts of multimedia data representations, color models, audio and video signals.", "K2"),
        ("CO2", "Apply the knowledge of various compression algorithms for developing multimedia applications.", "K3"),
        ("CO3", "Summarize the image compression standards, audio and video compression techniques.", "K2"),
        ("CO4", "Discuss the concepts of content-based image retrieval.", "K2"),
        ("CO5", "Describe the concept of cloud computing and its application in multimedia technologies.", "K2")
    ]
    
    for i, (co, desc, level) in enumerate(cos, 1):
        row = table.rows[i].cells
        row[0].text = co
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[1].text = desc
        row[2].text = level
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.add_run("Note: K1- Remember, K2- Understand, K3- Apply, K4- Analyze, K5- Evaluate, K6- Create")
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(9)
    
    doc.add_paragraph()
    
    # 5. CO-PO Mapping
    p = doc.add_paragraph()
    run = p.add_run("5. Mapping of Course Outcomes with Program Outcomes")
    run.bold = True
    run.font.size = Pt(11)
    
    table = create_table_with_borders(doc, 6, 12)
    
    header = table.rows[0].cells
    header[0].text = ""
    pos = ["PO1", "PO2", "PO3", "PO4", "PO5", "PO6", "PO7", "PO8", "PO9", "PO10", "PO11"]
    for i, po in enumerate(pos, 1):
        header[i].text = po
        header[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    mappings = [
        ["CO1", "3", "2", "1", "", "2", "", "", "", "", "", "2"],
        ["CO2", "3", "3", "2", "2", "2", "", "", "", "1", "", "2"],
        ["CO3", "3", "2", "2", "2", "3", "", "", "", "1", "", "2"],
        ["CO4", "2", "2", "2", "2", "2", "", "", "", "1", "", "2"],
        ["CO5", "2", "2", "2", "2", "3", "1", "1", "", "1", "1", "3"]
    ]
    
    for i, mapping in enumerate(mappings, 1):
        row = table.rows[i].cells
        for j, val in enumerate(mapping):
            row[j].text = val
            row[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 6. Text Books
    p = doc.add_paragraph()
    run = p.add_run("6. Text Books")
    run.bold = True
    run.font.size = Pt(11)
    
    table = create_table_with_borders(doc, 2, 5)
    
    header = table.rows[0].cells
    headers = ["Sl No", "Title of Book", "Name of Author/s", "Publisher", "Edition and Year"]
    for i, h in enumerate(headers):
        header[i].text = h
        for para in header[i].paragraphs:
            for run in para.runs:
                run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header[i], 'D9E2F3')
    
    row = table.rows[1].cells
    row[0].text = "1"
    row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row[1].text = "Fundamentals of Multimedia"
    row[2].text = "Ze-Nian Li and M. S. Drew"
    row[3].text = "Pearson Education"
    row[4].text = "2nd Edition, 2014"
    
    doc.add_paragraph()
    
    # 7. Reference Books
    p = doc.add_paragraph()
    run = p.add_run("7. Reference Books")
    run.bold = True
    run.font.size = Pt(11)
    
    table = create_table_with_borders(doc, 5, 5)
    
    header = table.rows[0].cells
    for i, h in enumerate(headers):
        header[i].text = h
        for para in header[i].paragraphs:
            for run in para.runs:
                run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header[i], 'D9E2F3')
    
    refs = [
        ("1", "Introduction to Multimedia Communications", "K. R. Rao, Zoran S. Bojkovic, D. A. Milovanovic", "Wiley", "1st Edition, 2006"),
        ("2", "Principles of Multimedia Database Systems", "V. S. Subrahmanian", "Morgan Kaufmann", "1998"),
        ("3", "Multimedia: Computing, Communications and Applications", "R. Steinmetz and K. Nahrstedt", "Pearson Education", "2002"),
        ("4", "Multimedia Systems", "John F. Koegel Buford", "Pearson Education", "1994")
    ]
    
    for i, ref in enumerate(refs, 1):
        row = table.rows[i].cells
        for j, val in enumerate(ref):
            row[j].text = val
            if j == 0:
                row[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 8. Video Links
    p = doc.add_paragraph()
    run = p.add_run("8. Video Links or any other Reference Materials")
    run.bold = True
    run.font.size = Pt(11)
    
    table = create_table_with_borders(doc, 4, 3)
    
    header = table.rows[0].cells
    header[0].text = "Sl No"
    header[1].text = "Details Like Weblink, Publication Details"
    header[2].text = "Module"
    for cell in header:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'D9E2F3')
    
    links = [
        ("1", "https://nptel.ac.in/courses/106105032 - Multimedia Systems", "1,2,3,4"),
        ("2", "https://nptel.ac.in/courses/117105082 - Digital Image Processing", "1,2,3"),
        ("3", "https://nptel.ac.in/courses/106106093 - Cloud Computing", "4")
    ]
    
    for i, link in enumerate(links, 1):
        row = table.rows[i].cells
        for j, val in enumerate(link):
            row[j].text = val
            if j in [0, 2]:
                row[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 9. Teaching Plan
    p = doc.add_paragraph()
    run = p.add_run("9. Teaching Plan")
    run.bold = True
    run.font.size = Pt(11)
    
    # Module 1 Teaching Plan
    table = create_table_with_borders(doc, 11, 3)
    
    header = table.rows[0].cells
    header[0].text = "No"
    header[1].text = "Topics"
    header[2].text = "No of hours\n(36 hrs)"
    for cell in header:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'D9E2F3')
    
    # Module 1 header
    row = table.rows[1].cells
    row[0].merge(row[1])
    row[0].text = "Module 1 (Multimedia Fundamentals & Data Representation)"
    for para in row[0].paragraphs:
        for run in para.runs:
            run.bold = True
    row[2].text = "9 hrs"
    for para in row[2].paragraphs:
        for run in para.runs:
            run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(row[0], 'E2EFDA')
    set_cell_shading(row[2], 'E2EFDA')
    
    m1_topics = [
        ("1.1", "Multimedia Basics- Multimedia, Hypermedia, WWW, Internet", "1"),
        ("1.2", "Multimedia Software- Editing and Authoring Tools", "1"),
        ("1.3", "Graphics/Image Data Types, Popular File Formats", "1"),
        ("1.4", "Color Science- Human visual system, Color perception", "1"),
        ("1.5", "Color Models- RGB, CMY/CMYK, HSV/HSL, YUV/YCbCr", "1"),
        ("1.6", "Digital Audio- Digitization of sound, Sampling, Quantization", "1"),
        ("1.7", "MIDI- Musical Instrument Digital Interface, Synthesis", "1"),
        ("1.8", "Digital Video- Video signals, Frame rates, Interlacing", "1"),
        ("1.9", "Video formats overview, Practical applications", "1")
    ]
    
    for i, (num, topic, hrs) in enumerate(m1_topics, 2):
        row = table.rows[i].cells
        row[0].text = num
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[1].text = topic
        row[2].text = hrs
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Module 2 Teaching Plan
    table = create_table_with_borders(doc, 11, 3)
    
    header = table.rows[0].cells
    header[0].text = "No"
    header[1].text = "Topics"
    header[2].text = "No of hours"
    for cell in header:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'D9E2F3')
    
    row = table.rows[1].cells
    row[0].merge(row[1])
    row[0].text = "Module 2 (Compression Algorithms & Techniques)"
    for para in row[0].paragraphs:
        for run in para.runs:
            run.bold = True
    row[2].text = "9 hrs"
    for para in row[2].paragraphs:
        for run in para.runs:
            run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(row[0], 'E2EFDA')
    set_cell_shading(row[2], 'E2EFDA')
    
    m2_topics = [
        ("2.1", "Information Theory Basics- Entropy, Redundancy, Compression fundamentals", "1"),
        ("2.2", "Run-Length Encoding (RLE)- Algorithm, Applications, Variants", "1"),
        ("2.3", "Variable-Length Coding- Huffman coding, Shannon-Fano coding", "1"),
        ("2.4", "Dictionary-Based Coding- LZ77, LZ78, LZW algorithms", "1"),
        ("2.5", "Arithmetic Coding- Principles, Encoder/Decoder, Adaptive methods", "1"),
        ("2.6", "Lossy Compression- Distortion measures (MSE, PSNR), Rate-Distortion theory", "1"),
        ("2.7", "Quantization Techniques- Scalar, Vector quantization, Lloyd-Max", "1"),
        ("2.8", "Transform Coding- DCT, DFT fundamentals", "1"),
        ("2.9", "Wavelet-Based Coding- Wavelet transforms, Multi-resolution analysis", "1")
    ]
    
    for i, (num, topic, hrs) in enumerate(m2_topics, 2):
        row = table.rows[i].cells
        row[0].text = num
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[1].text = topic
        row[2].text = hrs
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Module 3 Teaching Plan
    table = create_table_with_borders(doc, 11, 3)
    
    header = table.rows[0].cells
    header[0].text = "No"
    header[1].text = "Topics"
    header[2].text = "No of hours"
    for cell in header:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'D9E2F3')
    
    row = table.rows[1].cells
    row[0].merge(row[1])
    row[0].text = "Module 3 (Multimedia Compression Standards)"
    for para in row[0].paragraphs:
        for run in para.runs:
            run.bold = True
    row[2].text = "9 hrs"
    for para in row[2].paragraphs:
        for run in para.runs:
            run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(row[0], 'E2EFDA')
    set_cell_shading(row[2], 'E2EFDA')
    
    m3_topics = [
        ("3.1", "JPEG Standard- Baseline, DCT-based compression, Quality settings", "1"),
        ("3.2", "JPEG2000 & JPEG-LS- Wavelet-based, Lossless JPEG", "1"),
        ("3.3", "Bi-level Image Compression- JBIG, JBIG2, Fax standards", "1"),
        ("3.4", "Audio Compression- ADPCM in Speech Coding, Vocoders", "1"),
        ("3.5", "Psychoacoustic Models- Human auditory system, Masking effects", "1"),
        ("3.6", "MPEG Audio- MP3, AAC, Encoding/Decoding pipeline", "1"),
        ("3.7", "Video Compression Principles- Motion estimation, compensation", "1"),
        ("3.8", "MPEG-1 & MPEG-2- Video bitstream, I/P/B frames, Interlaced", "1"),
        ("3.9", "MPEG-4, MPEG-7 & H.264/AVC- Object-based, Metadata, Modern codecs", "1")
    ]
    
    for i, (num, topic, hrs) in enumerate(m3_topics, 2):
        row = table.rows[i].cells
        row[0].text = num
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[1].text = topic
        row[2].text = hrs
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Module 4 Teaching Plan
    table = create_table_with_borders(doc, 11, 3)
    
    header = table.rows[0].cells
    header[0].text = "No"
    header[1].text = "Topics"
    header[2].text = "No of hours"
    for cell in header:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'D9E2F3')
    
    row = table.rows[1].cells
    row[0].merge(row[1])
    row[0].text = "Module 4 (Advanced Applications & Cloud Multimedia)"
    for para in row[0].paragraphs:
        for run in para.runs:
            run.bold = True
    row[2].text = "9 hrs"
    for para in row[2].paragraphs:
        for run in para.runs:
            run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(row[0], 'E2EFDA')
    set_cell_shading(row[2], 'E2EFDA')
    
    m4_topics = [
        ("4.1", "Content-Based Image Retrieval (CBIR)- Feature extraction, Color histograms", "1"),
        ("4.2", "Similarity Measures & Indexing- Distance metrics, Query processing", "1"),
        ("4.3", "CBIR Case Study- CBIRD system, Practical implementation", "1"),
        ("4.4", "Video Retrieval & Search- Segmentation, Keyframe extraction, Querying", "1"),
        ("4.5", "Cloud Computing Overview- IaaS, PaaS, SaaS, Deployment models", "1"),
        ("4.6", "Multimedia Cloud Computing- Architecture, Components, Resource mgmt", "1"),
        ("4.7", "Cloud Media Services- Media sharing, Streaming, CDN integration", "1"),
        ("4.8", "Computation Offloading- Mobile cloud, Service partitioning, Video coding", "1"),
        ("4.9", "Interactive Cloud Gaming- Architecture, Latency, Quality of Experience", "1")
    ]
    
    for i, (num, topic, hrs) in enumerate(m4_topics, 2):
        row = table.rows[i].cells
        row[0].text = num
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[1].text = topic
        row[2].text = hrs
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # 10. Model Question Paper
    p = doc.add_paragraph()
    run = p.add_run("10. Model Question Paper")
    run.bold = True
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    # Question Paper Header
    p = doc.add_paragraph()
    run = p.add_run("MODEL QUESTION PAPER")
    run.bold = True
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run("APJ ABDUL KALAM TECHNOLOGICAL UNIVERSITY")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run("SIXTH SEMESTER B. TECH DEGREE EXAMINATION")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run("Course Code: CXT332")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run("Course Name: MULTIMEDIA TECHNOLOGIES")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    table = create_table_with_borders(doc, 1, 2)
    row = table.rows[0].cells
    row[0].text = "Max. Marks: 100"
    for para in row[0].paragraphs:
        for run in para.runs:
            run.bold = True
    row[1].text = "Duration: 3 hours"
    for para in row[1].paragraphs:
        for run in para.runs:
            run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    
    # Part A
    p = doc.add_paragraph()
    run = p.add_run("PART A")
    run.bold = True
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run("Answer all questions. Each question carries 3 marks")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Part A Questions Table
    table = create_table_with_borders(doc, 9, 4)
    
    header = table.rows[0].cells
    header[0].text = ""
    header[1].text = "Question"
    header[2].text = "CO"
    header[3].text = "Marks"
    for cell in header:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'D9E2F3')
    
    part_a = [
        ("1", "Differentiate between multimedia and hypermedia.", "1", "(3)"),
        ("2", "Explain the RGB color model with its applications.", "1", "(3)"),
        ("3", "What is the significance of entropy in data compression?", "2", "(3)"),
        ("4", "Compare Huffman coding and Arithmetic coding.", "2", "(3)"),
        ("5", "Describe the DCT transform used in JPEG.", "3", "(3)"),
        ("6", "What is psychoacoustic masking?", "3", "(3)"),
        ("7", "What is content-based image retrieval?", "4", "(3)"),
        ("8", "Define cloud service models (IaaS, PaaS, SaaS).", "4", "(3)")
    ]
    
    for i, (num, q, co, marks) in enumerate(part_a, 1):
        row = table.rows[i].cells
        row[0].text = num
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[1].text = q
        row[2].text = co
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[3].text = marks
        row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Part B
    p = doc.add_paragraph()
    run = p.add_run("PART B")
    run.bold = True
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run("Answer any one full question from each module. Each question carries 19 marks")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Part B Questions - 17 rows needed (1 header + 4 module headers + 12 questions)
    part_b = [
        ("", "", "Module 1", "", "", True),
        ("9", "a", "Explain various color models used in image processing.", "1", "(10)", False),
        ("", "b", "Describe the digitization process of audio signals.", "1", "(9)", False),
        ("10", "a", "Discuss different multimedia file formats and their characteristics.", "1", "(10)", False),
        ("", "b", "Explain the MIDI protocol and its applications.", "1", "(9)", False),
        ("", "", "Module 2", "", "", True),
        ("11", "a", "Explain Huffman coding with an example.", "2", "(10)", False),
        ("", "b", "Describe wavelet-based compression techniques.", "2", "(9)", False),
        ("12", "a", "Explain the LZW compression algorithm with an example.", "2", "(10)", False),
        ("", "b", "Describe vector quantization and its advantages.", "2", "(9)", False),
        ("", "", "Module 3", "", "", True),
        ("13", "a", "Describe the JPEG compression standard in detail.", "3", "(12)", False),
        ("", "b", "Compare JPEG and JPEG2000 standards.", "3", "(7)", False),
        ("14", "a", "Explain MPEG-1 video compression with block diagram.", "3", "(12)", False),
        ("", "b", "Describe MPEG audio compression.", "3", "(7)", False),
        ("", "", "Module 4", "", "", True),
        ("15", "a", "Explain content-based image retrieval techniques.", "4", "(10)", False),
        ("", "b", "Discuss computation offloading for multimedia services.", "4", "(9)", False)
    ]
    
    table = create_table_with_borders(doc, len(part_b) + 1, 5)
    
    header = table.rows[0].cells
    header[0].text = ""
    header[1].text = ""
    header[2].text = "Question"
    header[3].text = "CO"
    header[4].text = "Marks"
    for cell in header:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'D9E2F3')
    
    row_idx = 1
    for item in part_b:
        if item[5]:  # Module header
            row = table.rows[row_idx].cells
            row[0].merge(row[4])
            row[0].text = item[2]
            for para in row[0].paragraphs:
                for run in para.runs:
                    run.bold = True
            set_cell_shading(row[0], 'E2EFDA')
        else:
            row = table.rows[row_idx].cells
            row[0].text = item[0]
            row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for para in row[0].paragraphs:
                for run in para.runs:
                    run.bold = True
            row[1].text = item[1]
            row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for para in row[1].paragraphs:
                for run in para.runs:
                    run.bold = True
            row[2].text = item[2]
            row[3].text = item[3]
            row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[4].text = item[4]
            row[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_idx += 1
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Prepared By
    p = doc.add_paragraph()
    run = p.add_run("Prepared By:")
    run.bold = True
    
    doc.add_paragraph("Faculty Name: ____________________")
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Team members:")
    run.bold = True
    doc.add_paragraph("1. ____________________")
    doc.add_paragraph("2. ____________________")
    
    # Save document
    doc.save('/app/CXT332_Multimedia_Technologies_Syllabus_FINAL.docx')
    print("Document created successfully: CXT332_Multimedia_Technologies_Syllabus_FINAL.docx")

if __name__ == "__main__":
    main()
